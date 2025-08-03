from django.shortcuts import render
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status
from .serializers import UserRegistrationSerializer
from rest_framework.permissions import IsAuthenticated
from rest_framework.parsers import MultiPartParser, FormParser
from .models import File, User
from .serializers import FileSerializer
import chromadb
from .models import Chat, Message
from .serializers import ChatSerializer, MessageSerializer
from rest_framework import generics
import fitz  # PyMuPDF
from PIL import Image
import numpy as np
from io import BytesIO
from docx import Document
from langchain_openai import OpenAIEmbeddings as OE
from sentence_transformers import SentenceTransformer
from transformers import BlipProcessor, BlipForConditionalGeneration
import torch
import os
from tqdm import tqdm
import openai
from uuid import uuid4
import re
from pathlib import Path
from langchain_openai import OpenAIEmbeddings
from langchain.chains import RetrievalQA
from langchain.prompts import PromptTemplate
from langchain.text_splitter import RecursiveCharacterTextSplitter
try:
    from langchain_chroma import Chroma
except ImportError:
    from langchain_community.vectorstores import Chroma
from langchain_community.chat_models import ChatOpenAI
from langchain_google_genai import ChatGoogleGenerativeAI
from dotenv import load_dotenv
from django.shortcuts import get_object_or_404
import sqlite3
import subprocess

load_dotenv()

from django.conf import settings
from django.db import transaction, models
from django.utils import timezone

# Initialize ChromaDB clients (simple local setup)
chroma_client = chromadb.Client(chromadb.config.Settings(
    persist_directory="chroma_data"
))

# Load BLIP for image captioning
processor = BlipProcessor.from_pretrained("Salesforce/blip-image-captioning-base", use_fast=True)
blip_model = BlipForConditionalGeneration.from_pretrained("Salesforce/blip-image-captioning-base").eval()
device = "cuda" if torch.cuda.is_available() else "cpu"
blip_model.to(device)

text_emb = OE(model="text-embedding-3-small")
#image_emb = SentenceTransformer("clip-ViT-B-32")

# --- PDF and Image Processing ---
def describe_image(image_path):
    raw_image = Image.open(image_path).convert('RGB')
    try:
        inputs = processor([raw_image], padding=True, return_tensors="pt").to(device)
        out = blip_model.generate(**inputs)
        # Check if output tensor has at least one sequence
        if out is None or out.shape[0] == 0:
            return "No caption generated"
        caption = processor.decode(out[0], skip_special_tokens=True)
        return caption
    except Exception as e:
        return "No caption generated"

def extract_pdf_pages(pdf_path, image_dir="imgs"):
    os.makedirs(image_dir, exist_ok=True)
    doc = fitz.open(pdf_path)
    pages = []
    for i, page in enumerate(doc, 1):
        txt = page.get_text()
        imgs = []
        for img in page.get_images(full=True):
            xref = img[0]
            try:
                pix = fitz.Pixmap(doc, xref)
                if pix.colorspace.n not in [1, 3]:
                    pix = fitz.Pixmap(fitz.csRGB, pix)
                buf = pix.tobytes("png")
                path = f"{image_dir}/pg{i}_img{xref}.png"
                Image.open(BytesIO(buf)).save(path)
                imgs.append(path)
            except Exception:
                pass
        pages.append({"page": i, "text": txt, "images": imgs})
    return pages

def extract_docx_text(docx_path):
    """Extract text from DOCX file"""
    try:
        doc = Document(docx_path)
        full_text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text)
        
        return {"page": 1, "text": "\n".join(full_text), "images": []}
    except Exception as e:
        raise ValueError(f"Failed to extract text from DOCX file: {str(e)}")

def build_embeddings(pages, text_batch_size=50):
    texts, metadatas = [], []
    for p in pages:
        if p["text"].strip():
            texts.append(p["text"])
            metadatas.append({"type": "text", "page": p["page"]})
        for im in p["images"]:
            caption = describe_image(im)
            if not caption or caption.startswith("Image captioning failed"):
                caption = "No caption generated"
            texts.append(caption)
            metadatas.append({
                "type": "image_caption",
                "page": p["page"],
                "path": im,
                "caption": caption
            })
    text_vecs = []
    if texts:
        for i in range(0, len(texts), text_batch_size):
            batch = texts[i : i + text_batch_size]
            text_vecs.extend(text_emb.embed_documents(batch))  # Only use OpenAIEmbeddings
    return text_vecs, texts, metadatas

# --- Per-user Chroma directory helper ---
def get_user_chroma_dir(user_or_id):
    if isinstance(user_or_id, int) or isinstance(user_or_id, str):
        user_id = user_or_id
    else:
        user_id = user_or_id.id
    return os.path.join("chroma_data", f"user_{user_id}")

# --- File upload: extract, caption, embed, and store in per-user Chroma ---
def process_and_store_file(user, file_path, collection_name, file_id=None):
    try:
        print(f"Processing file: {file_path}")
        
        # Check if file exists
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Determine file type and extract content accordingly
        file_extension = file_path.lower().split('.')[-1]
        print(f"File extension: {file_extension}")
        
        if file_extension == 'pdf':
            print("Extracting PDF pages...")
            pages = extract_pdf_pages(file_path)
        elif file_extension == 'docx':
            print("Extracting DOCX text...")
            # For DOCX, we get a single page with all text
            pages = [extract_docx_text(file_path)]
        else:
            raise ValueError(f"Unsupported file type: {file_extension}. Only PDF and DOCX are supported.")
        
        print(f"Extracted {len(pages)} pages")
        
        texts, metas, ids = [], [], []
        splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
        chunk_idx = 0
        
        for p in pages:
            if p["text"].strip():
                chunks = splitter.split_text(p["text"])
                for chunk in chunks:
                    texts.append(chunk)
                    metas.append({"type": "text", "page": p["page"], "file_id": file_id})
                    ids.append(f"{file_id}_{chunk_idx}")
                    chunk_idx += 1
            for im in p["images"]:
                caption = describe_image(im)
                texts.append(caption)
                metas.append({
                    "type": "image_caption",
                    "page": p["page"],
                    "file_id": file_id,
                    "path": im,
                    "caption": caption
                })
                ids.append(f"{file_id}_img_{chunk_idx}")
                chunk_idx += 1

        print(f"Created {len(texts)} text chunks")
        print(f"DEBUG: First few metadata entries: {metas[:3] if metas else 'No metadata'}")
        print(f"DEBUG: First few IDs: {ids[:3] if ids else 'No IDs'}")
        
        if not texts:
            raise ValueError(f"{file_extension.upper()} file contained no text or images.")
        
        user_chroma_dir = get_user_chroma_dir(user)
        print(f"Chroma directory: {user_chroma_dir}")
        
        # Check and fix permissions before proceeding
        if not check_and_fix_chroma_permissions(user_chroma_dir):
            raise Exception(f"Failed to set proper permissions for ChromaDB directory: {user_chroma_dir}")
        
        # Ensure directory exists
        os.makedirs(user_chroma_dir, exist_ok=True)
        
        vs = Chroma(
            collection_name=collection_name,
            embedding_function=text_emb,
            persist_directory=user_chroma_dir,
        )
        
        print(f"Adding {len(texts)} texts to ChromaDB...")
        batch_size = 100
        for i in range(0, len(texts), batch_size):
            batch_texts = texts[i:i+batch_size]
            batch_metas = metas[i:i+batch_size]
            batch_ids = ids[i:i+batch_size]
            try:
                vs.add_texts(
                    texts=batch_texts,
                    metadatas=batch_metas,
                    ids=batch_ids
                )
                print(f"Added batch {i//batch_size + 1}/{(len(texts) + batch_size - 1)//batch_size}")
            except Exception as batch_error:
                print(f"❌ Error adding batch to ChromaDB: {batch_error}")
                if "readonly database" in str(batch_error).lower():
                    print(f"🔍 ChromaDB database is read-only. Checking permissions...")
                    import stat
                    if os.path.exists(user_chroma_dir):
                        st = os.stat(user_chroma_dir)
                        print(f"🔍 Directory permissions: {oct(st.st_mode)}")
                        print(f"🔍 Directory owner: {st.st_uid}")
                        print(f"🔍 Directory group: {st.st_gid}")
                        print(f"🔍 Current user: {os.getuid()}")
                        print(f"🔍 Current group: {os.getgid()}")
                    raise Exception(f"ChromaDB database is read-only. Please check file permissions for {user_chroma_dir}")
                else:
                    raise batch_error
        
        # The newer version of langchain-chroma doesn't have persist() method
        # The data is automatically persisted when using persist_directory
        print(f"Successfully stored {len(texts)} documents for user {user.id} in {user_chroma_dir}")
        
        # Clean up extracted images after processing
        for p in pages:
            for im in p["images"]:
                try:
                    os.remove(im)
                except Exception:
                    pass  # Ignore errors during cleanup
                    
    except Exception as e:
        print(f"Error in process_and_store_file: {str(e)}")
        import traceback
        traceback.print_exc()
        raise e

# --- RAG query: load per-user Chroma and run RetrievalQA ---
def run_rag_query(user, query, collection_name, llm=None, k=5):
    user_chroma_dir = get_user_chroma_dir(user)
    if not Path(user_chroma_dir).exists():
        # For general questions, provide a helpful response without mentioning documents
        if any(word in query.lower() for word in ['hello', 'hi', 'hey', 'greetings', 'how are you', 'good morning', 'good afternoon', 'good evening']):
            return "Hello! I'm here to help you with your questions. Feel free to ask me anything!", []
        return "I don't have any documents to reference, but I'm happy to help with general questions!", []
    
    vs = Chroma(
        persist_directory=user_chroma_dir,
        embedding_function=text_emb,
        collection_name=collection_name,
    )
    prompt_template = (
        "You are a helpful assistant. Use the provided context and conversation history to answer the question. "
        "If the question refers to previous messages in the conversation, make sure to consider that context. "
        "When you find relevant information in the context, provide specific details and facts from the documents. "
        "If you cannot find specific information in the context, respond with a helpful message like 'I don't have enough information to answer this question' or 'I cannot find specific information about that in the available documents.' "
        "Be honest about what information is available and what is not. Provide detailed answers when you have specific information from the documents.\n\n"
        "Context:\n{context}\n\nQuestion: {question}\nAnswer:"
    )
    prompt = PromptTemplate(
        template=prompt_template,
        input_variables=["context", "question"]
    )
    if llm is None:
        llm = ChatOpenAI(model_name="gpt-4.1-mini", openai_api_key=settings.OPENAI_API_KEY, temperature=0.2)
    rag_chain = RetrievalQA.from_chain_type(
        llm=llm,
        retriever=vs.as_retriever(search_kwargs={"k": k}),
        return_source_documents=True,
        chain_type_kwargs={"prompt": prompt}
    )
    try:
        result = rag_chain.invoke({"query": query})
        answer = result["result"]
        sources = result.get("source_documents", [])
        
        # Check if sources are actually relevant (have reasonable similarity scores)
        # This is a heuristic - if the answer is very generic, the sources might not be relevant
        print(f"Retrieved {len(sources)} source documents")
        
        # Simple logic: only include sources if documents were actually retrieved
        if not sources:
            print(f"No documents retrieved - returning empty sources")
            return answer, []
        
        # If documents were retrieved, include them as sources
        print(f"Documents retrieved - including sources")
        return answer, sources
    except Exception as e:
        print(f"RAG query failed: {str(e)}")
        return "I apologize, but I encountered an error while processing your request. Please try again or contact support if the issue persists.", []


def run_rag_pipeline(user, query, chat_id=None):
    # Check for document listing request first
    query_lower = query.lower().strip()
    document_list_keywords = [
        'what documents', 'what are the documents', 'what files', 'what are the files',
        'list documents', 'list files', 'show documents', 'show files',
        'what documents do i have', 'what files do i have', 'my documents', 'my files',
        'uploaded documents', 'uploaded files', 'documents i have', 'files i have',
        'which documents', 'which files', 'tell me which documents', 'tell me which files',
        'can you tell me which documents', 'can you tell me which files',
        'what documents did i upload', 'what files did i upload',
        'which documents did i upload', 'which files did i upload',
        'show me my documents', 'show me my files', 'display my documents', 'display my files',
        'what have i uploaded', 'what did i upload', 'my uploads', 'my uploaded files',
        'can you list', 'list me', 'show me the', 'tell me the', 'what docs', 'what files',
        'my docs', 'my files', 'documents list', 'files list', 'uploaded docs', 'uploaded files',
        'available documents', 'available files', 'stored documents', 'stored files',
        'give me the list', 'give me list', 'show me list', 'tell me list',
        'what documents have i', 'what files have i', 'my document list', 'my file list',
        'list of documents', 'list of files', 'documents list', 'files list',
        'how many documents', 'how many files', 'count documents', 'count files'
    ]
    
    is_document_list_request = any(keyword in query_lower for keyword in document_list_keywords)
    print(f"Query: '{query_lower}'")
    print(f"Is document list request: {is_document_list_request}")
    if is_document_list_request:
        print("✅ Document list request detected!")
        print(f"🔍 Keywords matched: {[keyword for keyword in document_list_keywords if keyword in query_lower]}")
    
    # Get conversation context if chat_id is provided
    conversation_context = ""
    if chat_id:
        try:
            # Get recent messages from the chat (last 10 messages for context)
            recent_messages = Message.objects.filter(
                chat_id=chat_id,
                chat__user_id=user.id
            ).order_by('-timestamp')[:10]
            
            # Build conversation context from recent messages
            if recent_messages:
                conversation_parts = []
                for msg in reversed(recent_messages):  # Reverse to get chronological order
                    if msg.sender_type == "user":
                        conversation_parts.append(f"User: {msg.content}")
                    elif msg.sender_type == "ai":
                        # Include sources in AI responses for better context only if sources exist
                        ai_response = msg.content
                        if hasattr(msg, 'sources') and msg.sources and len(msg.sources) > 0:
                            ai_response += f"\n[Source: 1 document referenced]"
                        conversation_parts.append(f"Assistant: {ai_response}")
                
                conversation_context = "\n".join(conversation_parts)
                print(f"Conversation context length: {len(conversation_context)} characters")
        except Exception as e:
            print(f"Error getting conversation context: {str(e)}")

    # Handle document listing request
    if is_document_list_request:
        try:
            print("🔍 Using ChromaDB query for document listing...")
            print(f"🔍 Current user in session: {user.id}")
            
            # Test the ChromaDB document listing for debugging
            test_user_chroma_document_listing(user)
            
            # Use the ChromaDB approach to get files that have been processed
            file_list = get_user_files_from_chroma(user)
            
            if file_list:
                print(f"✅ Found {len(file_list)} files in database for user {user.id}")
                
                # Create response with files found in ChromaDB
                if len(file_list) == 1:
                    response_text = f"📚 You have 1 document available for AI queries:\n\n"
                else:
                    response_text = f"📚 You have {len(file_list)} documents available for AI queries:\n\n"
                
                for i, file_info in enumerate(file_list, 1):
                    # Create a detailed summary based on file type and name
                    filename = file_info['filename']
                    file_type = file_info['file_type']
                    upload_date = file_info['uploaded_at']
                    
                    # Generate a detailed summary based on filename analysis
                    filename_lower = filename.lower()
                    
                    if 'proposal' in filename_lower:
                        if 'fine' in filename_lower and 'tuning' in filename_lower:
                            summary = "This appears to be a proposal document about fine-tuning AI systems, likely containing technical specifications and implementation details for AI model optimization"
                        else:
                            summary = "This appears to be a proposal document, typically containing project plans, objectives, and implementation strategies"
                    elif 'counterclaim' in filename_lower or 'claim' in filename_lower:
                        summary = "This appears to be a legal document containing counterclaims, likely part of legal proceedings or dispute resolution"
                    elif 'report' in filename_lower:
                        summary = "This appears to be a report document, typically containing analysis, findings, and recommendations"
                    elif 'contract' in filename_lower or 'agreement' in filename_lower:
                        summary = "This appears to be a contract or agreement document, containing legal terms and conditions"
                    elif 'manual' in filename_lower or 'guide' in filename_lower:
                        summary = "This appears to be a manual or guide document, containing instructions and procedures"
                    elif 'invoice' in filename_lower or 'bill' in filename_lower:
                        summary = "This appears to be an invoice or billing document, containing financial transaction details"
                    elif 'resume' in filename_lower or 'cv' in filename_lower:
                        summary = "This appears to be a resume or CV document, containing professional background and qualifications"
                    elif 'presentation' in filename_lower or 'ppt' in filename_lower:
                        summary = "This appears to be a presentation document, likely containing slides and visual content"
                    elif 'spreadsheet' in filename_lower or 'excel' in filename_lower or 'csv' in filename_lower:
                        summary = "This appears to be a spreadsheet or data document, containing structured data and calculations"
                    elif 'email' in filename_lower or 'pst' in filename_lower:
                        summary = "This appears to be an email archive or communication document, containing email correspondence"
                    elif 'image' in filename_lower or 'photo' in filename_lower or 'jpeg' in filename_lower or 'png' in filename_lower:
                        summary = "This appears to be an image or visual document, containing graphical content"
                    else:
                        summary = f"This is a {file_type} document that may contain various types of content"
                    
                    # Add upload date for context with better formatting
                    response_text += f"📄 **{filename}**\n   {summary}\n   📅 Uploaded: {upload_date}\n\n"
                
                response_text += f"\n💡 You can ask me questions about any of these documents! For example:\n• 'What is the main topic of {file_list[0]['filename'] if file_list else 'your document'}?'\n• 'Summarize the key points from my documents'\n• 'Find information about [specific topic] in my documents'"
                
                # Return with sources showing all available documents
                sources = []
                for file_info in file_list:
                    sources.append({
                        "content": f"Document available for queries: {file_info['filename']}",
                        "page": "1",
                        "file_id": str(file_info['file_id']),
                        "type": "document_list",
                        "filename": file_info['filename']
                    })
                
                return {"answer": response_text, "sources": sources}
            else:
                return {"answer": "📭 You haven't uploaded any documents yet.\n\n💡 Try uploading a document and then ask about your documents again.", "sources": []}
        except Exception as e:
            print(f"Error listing documents: {str(e)}")
            return {"answer": "❌ I encountered an error while retrieving your document list. Please try again or contact support if the issue persists.", "sources": []}

    # First, try to retrieve documents to assess confidence
    collection_name = get_chroma_collection_name(user)
    user_chroma_dir = get_user_chroma_dir(user)
    
    # Check if user has any documents
    if not Path(user_chroma_dir).exists():
        print(f"No documents found for user - using conversation context only")
        # Use conversation context for response
        if conversation_context:
            context_prompt = f"Based on the following conversation history, answer the question naturally:\n\nConversation History:\n{conversation_context}\n\nQuestion: {query}\nAnswer:"
            try:
                # Use OpenAI for consistency
                llm_instance = ChatOpenAI(model_name="gpt-4.1-mini", openai_api_key=settings.OPENAI_API_KEY, temperature=0.2)
                response = llm_instance.invoke(context_prompt)
                answer = response.content if hasattr(response, 'content') else str(response)
                return {"answer": answer, "sources": []}
            except Exception as e:
                return {"answer": "I don't have any documents to reference, but I'm happy to help with general questions!", "sources": []}
        else:
            return {"answer": "I don't have any documents to reference, but I'm happy to help with general questions!", "sources": []}
    
    # Try to retrieve documents
    try:
        vs = Chroma(
            persist_directory=user_chroma_dir,
            embedding_function=text_emb,
            collection_name=collection_name,
        )
        
        # Get document similarity scores
        results = vs.similarity_search_with_score(query, k=5)
        print(f"Retrieved {len(results)} documents with scores")
        
        # Calculate confidence score
        if results:
            scores = [score for _, score in results]
            max_confidence = min(scores)  # Lower distance = higher confidence
            print(f"Max confidence: {max_confidence}")
            
            # Always try to use document retrieval if we have any results, regardless of confidence
            if results:
                print(f"Using document retrieval with confidence {max_confidence:.3f}")
                
                # Use full RAG pipeline with documents
                answer, sources = run_rag_query(user, query, collection_name)
                
                # Process sources - always include the best available document
                enhanced_sources = []
                if sources and len(sources) > 0:
                    best_doc = sources[0]
                    print(f"Retrieved {len(sources)} documents, using highest match as source")
                    
                    if hasattr(best_doc, 'metadata') and best_doc.metadata:
                        source_info = {
                            "content": best_doc.page_content,
                            "page": best_doc.metadata.get("page", "Unknown"),
                            "file_id": best_doc.metadata.get("file_id", "Unknown"),
                            "type": best_doc.metadata.get("type", "text")
                        }
                        
                        try:
                            if source_info["file_id"] != "Unknown":
                                file_obj = File.objects.get(id=source_info["file_id"])
                                source_info["filename"] = file_obj.filename
                            else:
                                source_info["filename"] = "Unknown"
                        except File.DoesNotExist:
                            source_info["filename"] = "Unknown"
                            
                        enhanced_sources.append(source_info)
                    else:
                        enhanced_sources.append({
                            "content": best_doc.page_content,
                            "page": "Unknown",
                            "file_id": "Unknown",
                            "type": "text",
                            "filename": "Unknown"
                        })
                    
                # If no sources from RAG, use the best document from similarity search
                if not enhanced_sources and results:
                    best_result = results[0]  # Get the highest confidence document
                    best_doc_content, best_score = best_result
                    print(f"Using best similarity result with confidence {best_score:.3f}")
                    
                    # Create a source from the best similarity result
                    enhanced_sources.append({
                        "content": best_doc_content.page_content,
                        "page": "1",
                        "file_id": "similarity_match",
                        "type": "text",
                        "filename": "Best matching document"
                    })
                
                return {"answer": answer, "sources": enhanced_sources}
            else:
                print("No documents retrieved - returning no information message")
                return {"answer": "I'm not able to find information from your uploaded documents that answers your question. Please try rephrasing your question or ask about a different topic.", "sources": []}
        else:
            print("No documents retrieved - returning no information message")
            return {"answer": "I'm not able to find information from your uploaded documents that answers your question. Please try rephrasing your question or ask about a different topic.", "sources": []}
    except Exception as e:
        print(f"Error in document retrieval: {str(e)}")
        return {"answer": "I'm not able to find information from your uploaded documents that answers your question. Please try rephrasing your question or ask about a different topic.", "sources": []}


from rest_framework.permissions import IsAuthenticated
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status

class UserLLMModelView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request):
        user = request.user
        selected = getattr(user, "preferred_llm", "openai")
        
        # Get additional information about the model
        model_info = {
            "preferred_llm": selected,
            "available_models": ["openai", "gemini"],
            "current_model_details": {
                "name": selected,
                "provider": "OpenAI" if selected == "openai" else "Google",
                "model_id": "gpt-4.1-mini" if selected == "openai" else "gemini-2.5-flash-lite",
                "status": "active"
            }
        }
        
        return Response(model_info, status=status.HTTP_200_OK)

    def post(self, request):
        user = request.user
        new_model = request.data.get("preferred_llm")
        if new_model not in ("openai", "gemini"):
            return Response({"error": "Invalid LLM model. Choose 'openai' or 'gemini'."}, status=status.HTTP_400_BAD_REQUEST)
        setattr(user, "preferred_llm", new_model)
        user.save()
        return Response({"message": f"Preferred LLM model set to '{new_model}'"}, status=status.HTTP_200_OK)

def get_chroma_collection_name(user):
    username = re.sub(r'[^a-zA-Z0-9._-]', '_', user.username)
    if 3 <= len(username) <= 512 and username[0].isalnum() and username[-1].isalnum():
        return username
    return f"user_{user.id}"

def clean_chroma_sqlite(user_chroma_dir, collection_name):
    """
    Clean ChromaDB SQLite database by removing entries for specific collection
    """
    try:
        sqlite_path = os.path.join(user_chroma_dir, "chroma.sqlite3")
        if not os.path.exists(sqlite_path):
            print(f"ChromaDB SQLite file not found: {sqlite_path}")
            return False
            
        print(f"Cleaning ChromaDB SQLite: {sqlite_path}")
        
        # Connect to SQLite database
        conn = sqlite3.connect(sqlite_path)
        cursor = conn.cursor()
        
        # Get all tables
        cursor.execute("SELECT name FROM sqlite_master WHERE type='table';")
        tables = cursor.fetchall()
        print(f"Found tables: {[table[0] for table in tables]}")
        
        # Check embeddings table structure
        cursor.execute("PRAGMA table_info(embeddings);")
        columns = cursor.fetchall()
        print(f"Embeddings table columns: {[col[1] for col in columns]}")
        
        # Get total count before deletion
        cursor.execute("SELECT COUNT(*) FROM embeddings;")
        total_before = cursor.fetchone()[0]
        print(f"Total embeddings before deletion: {total_before}")
        
        # Show some sample data with correct column names
        try:
            cursor.execute("SELECT id, embedding_id, seq_id FROM embeddings LIMIT 5;")
            samples = cursor.fetchall()
            print(f"Sample embeddings: {samples}")
        except Exception as e:
            print(f"Could not fetch sample data: {e}")
        
        # Get segment_id for the collection first
        cursor.execute("SELECT id FROM segments WHERE name = ?;", (collection_name,))
        segment_result = cursor.fetchone()
        
        if segment_result:
            segment_id = segment_result[0]
            print(f"Found segment_id: {segment_id} for collection: {collection_name}")
            
            # Delete embeddings for the specific segment/collection
            cursor.execute("DELETE FROM embeddings WHERE segment_id = ?;", (segment_id,))
            deleted_count = cursor.rowcount
            
            # Get count after deletion
            cursor.execute("SELECT COUNT(*) FROM embeddings;")
            total_after = cursor.fetchone()[0]
            
            print(f"Deleted embeddings: {deleted_count}")
            print(f"Total embeddings after deletion: {total_after}")
            
            # Also clean up the segment metadata
            cursor.execute("DELETE FROM segment_metadata WHERE segment_id = ?;", (segment_id,))
            segment_metadata_deleted = cursor.rowcount
            print(f"Deleted segment metadata entries: {segment_metadata_deleted}")
            
            # Clean up embedding metadata
            cursor.execute("DELETE FROM embedding_metadata WHERE segment_id = ?;", (segment_id,))
            embedding_metadata_deleted = cursor.rowcount
            print(f"Deleted embedding metadata entries: {embedding_metadata_deleted}")
            
        else:
            print(f"No segment found for collection: {collection_name}")
            # Fallback: try to delete by embedding_id pattern
            cursor.execute("DELETE FROM embeddings WHERE embedding_id LIKE ?;", (f"{collection_name}%",))
            deleted_count = cursor.rowcount
            print(f"Fallback deletion: {deleted_count} embeddings")
        
        # Commit changes
        conn.commit()
        conn.close()
        
        return True
        
    except Exception as e:
        print(f"Error cleaning ChromaDB SQLite: {e}")
        return False

# Create your views here.

from supabase import create_client, Client
from supabase.lib.client_options import ClientOptions
from gotrue.errors import AuthApiError

from datetime import datetime


# SUPABASE_URL = os.getenv('SUPABASE_URL')
# SUPABASE_KEY = os.getenv('SUPABASE_KEY')
supabase: Client = create_client(settings.SUPABASE_URL, settings.SUPABASE_KEY)
class RegisterView(APIView):
    def post(self, request):
        data = request.data
        required = [
            'email', 'password',
            'first_name','last_name',
            'phone_number','gender','date_of_birth'
        ]
        missing = [f for f in required if not data.get(f)]
        if missing:
            return Response(
                {'error': f"Missing fields: {', '.join(missing)}"},
                status=status.HTTP_400_BAD_REQUEST
            )

        supabase = create_client(
            settings.SUPABASE_URL,
            settings.SUPABASE_SERVICE_ROLE_KEY
        )


        try:
            resp = supabase.auth.sign_up({
                'email':    data['email'],
                'password': data['password'],
                'options': {
                    'redirect_to': settings.BASE_URL_SIGNIN
                }
            })
        except AuthApiError as e:
            # If user already exists but is not verified, resend verification email
            if "User already registered" in str(e) or "already been registered" in str(e):
                try:
                    # Resend verification email for existing unverified user
                    resend_resp = supabase.auth.resend({
                        'type': 'signup',
                        'email': data['email'],
                        'options': {
                            'redirect_to': settings.BASE_URL_SIGNIN
                        }
                    })
                    return Response(
                        {'message': 'Verification email has been resent. Please check your email and verify your account.'},
                        status=status.HTTP_200_OK
                    )
                except Exception as resend_error:
                    return Response(
                        {'error': f'Failed to resend verification email: {str(resend_error)}'},
                        status=status.HTTP_500_INTERNAL_SERVER_ERROR
                    )
            else:
                return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)

        # At this point, resp.user is guaranteed to exist
        supa_user = resp.user  # <— has .id and .email :contentReference[oaicite:0]{index=0}

        # Push metadata into Supabase
        metadata = {
            'first_name':   data['first_name'],
            'last_name':    data['last_name'],
            'phone_number': data['phone_number'],
            'gender':       data['gender'],
            'date_of_birth': data['date_of_birth']
        }
        try:
            supabase.auth.admin.update_user_by_id(
                supa_user.id,
                {'user_metadata': metadata}
            )
        except AuthApiError as e:
            return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

        # Mirror into Django
        try:
            dob = datetime.strptime(data['date_of_birth'], "%Y-%m-%d").date()
        except ValueError:
            dob = None

        User.objects.create_user(
            id=supa_user.id,
            email=supa_user.email,
            first_name=metadata['first_name'],
            last_name=metadata['last_name'],
            phone_number=metadata['phone_number'],
            gender=metadata['gender'],
            date_of_birth=dob
        )

        return Response(
            {'message': 'Registration successful. Please verify your email.'},
            status=status.HTTP_201_CREATED
        )


from rest_framework import generics
from rest_framework.permissions import IsAuthenticated
from .serializers import UserSerializer

class UserProfileView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request, *args, **kwargs):
        supabase = create_client(settings.SUPABASE_URL, settings.SUPABASE_SERVICE_ROLE_KEY)
        user = request.user

        # Fetch fresh user data from Supabase using the user's UUID
        response = supabase.auth.admin.get_user_by_id(user.id)
        if response.get('error'):
            return Response({"error": response['error']['message']}, status=status.HTTP_400_BAD_REQUEST)

        raw_user_data = response.get('data')
        if not raw_user_data:
            return Response({"error": "User not found in Supabase."}, status=status.HTTP_404_NOT_FOUND)

        # Construct profile picture URL if available
        profile_picture_key = raw_user_data.get('user_metadata', {}).get('profile_picture')
        profile_picture_url = None
        bucket_name = "user-uploads"
        if profile_picture_key:
            if profile_picture_key.startswith("http"):
                profile_picture_url = profile_picture_key
            else:
                profile_picture_url = f"{settings.SUPABASE_URL}/storage/v1/object/public/{bucket_name}/{profile_picture_key}"

        # Compose all user data to return
        user_metadata = raw_user_data.get('user_metadata', {})
        bucket_name = "user-uploads"
        if profile_picture_key and not profile_picture_key.startswith("http"):
            profile_picture_url = f"{settings.SUPABASE_URL}/storage/v1/object/public/{bucket_name}/{profile_picture_key}"
            user_metadata['profile_picture'] = profile_picture_url
        elif profile_picture_key and profile_picture_key.startswith("http"):
            user_metadata['profile_picture'] = profile_picture_key
        else:
            user_metadata['profile_picture'] = None

        result = {
            "id": raw_user_data.get('id'),
            "email": raw_user_data.get('email'),
            "phone_confirmed_at": raw_user_data.get('phone_confirmed_at'),
            "email_confirmed_at": raw_user_data.get('email_confirmed_at'),
            "last_sign_in_at": raw_user_data.get('last_sign_in_at'),
            "app_metadata": raw_user_data.get('app_metadata'),
            "user_metadata": user_metadata,
            "created_at": raw_user_data.get('created_at'),
            "updated_at": raw_user_data.get('updated_at')
        }

        return Response(result, status=status.HTTP_200_OK)

    def put(self, request, *args, **kwargs):
        user = request.user
        data = request.data
        supabase = create_client(settings.SUPABASE_URL, settings.SUPABASE_SERVICE_ROLE_KEY)

        user_metadata = {
            "first_name": data.get('first_name', ''),
            "last_name": data.get('last_name', ''),
            "phone_number": data.get('phone_number', ''),
            "gender": data.get('gender', ''),
            "date_of_birth": data.get('date_of_birth', ''),
            "profile_picture": data.get('profile_picture', '')
        }

        updates = {
            "user_metadata": user_metadata
        }

        response = supabase.auth.admin.update_user_by_id(user.id, updates)
        if response.get('error'):
            return Response({"error": response['error']['message']}, status=status.HTTP_400_BAD_REQUEST)

        # Also update local Django user model if exists
        local_user = user
        local_user.first_name = user_metadata.get('first_name')
        local_user.last_name = user_metadata.get('last_name')
        local_user.phone_number = user_metadata.get('phone_number')
        local_user.gender = user_metadata.get('gender')
        if user_metadata.get('date_of_birth'):
            from datetime import datetime
            try:
                local_user.date_of_birth = datetime.strptime(user_metadata.get('date_of_birth'), "%Y-%m-%d").date()
            except Exception:
                local_user.date_of_birth = None
        local_user.profile_picture = user_metadata.get('profile_picture')
        local_user.save()

        return Response({"message": "User profile updated successfully."}, status=status.HTTP_200_OK)



from rest_framework.parsers import MultiPartParser, FormParser
from rest_framework.response import Response
from rest_framework import status


class UserProfilePictureGetView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request, *args, **kwargs):
        user = request.user
        profile_picture_key = user.profile_picture
        if not profile_picture_key:
            return Response({"error": "User has no profile picture."}, status=status.HTTP_404_NOT_FOUND)

        supabase_url = settings.SUPABASE_URL
        bucket_name = "user-uploads"

        # Construct public URL (assuming bucket is public or using signed URLs would require more implementation)
        image_url = f"{supabase_url}/storage/v1/object/public/{bucket_name}/{profile_picture_key}"

        return Response({"profile_picture_url": image_url}, status=status.HTTP_200_OK)


class UserProfilePictureUploadView(APIView):
    permission_classes = [IsAuthenticated]
    parser_classes = [MultiPartParser, FormParser]

    def post(self, request, *args, **kwargs):
        user = request.user
        file_obj = request.FILES.get('profile_picture')
        if not file_obj:
            return Response({"error": "No profile_picture file provided."}, status=status.HTTP_400_BAD_REQUEST)

        folder = f"user_{user.id}"
        filename = file_obj.name
        storage_key = f"{folder}/{filename}"

        supabase = create_client(settings.SUPABASE_URL, settings.SUPABASE_SERVICE_ROLE_KEY)

        try:
            file_bytes = file_obj.read()

            # Upload to Supabase Storage
            res = supabase.storage.from_("user-uploads").upload(
                storage_key,
                file_bytes,
                {"content-type": file_obj.content_type}
            )
            if isinstance(res, dict) and res.get("error"):
                return Response({"error": res["error"]["message"]}, status=status.HTTP_400_BAD_REQUEST)

            # Set profile_picture field to the storage key or URL
            user.profile_picture = storage_key  # or construct a full URL if needed
            user.save(update_fields=['profile_picture'])

            # No local file to delete since file is read directly from upload

            serializer = UserSerializer(user)
            return Response(serializer.data, status=status.HTTP_200_OK)

        except Exception as e:
            return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


class LoginView(APIView):
    def post(self, request):
        data = request.data
        email = data.get('email')
        password = data.get('password')

        if not email or not password:
            return Response({'error': 'Email and password are required.'}, status=status.HTTP_400_BAD_REQUEST)

        try:
            response = supabase.auth.sign_in_with_password(
                {'email': email, 'password': password}
            )
        except Exception as e:
            return Response({'error': 'Failed to login. ' + str(e)}, status=status.HTTP_400_BAD_REQUEST)

        if hasattr(response, 'error') and response.error:
            return Response({'error': str(response.error)}, status=status.HTTP_400_BAD_REQUEST)

        user = response.user
        if not user:
            return Response({'error': 'Invalid login credentials.'}, status=status.HTTP_401_UNAUTHORIZED)

        return Response({
            'access_token': response.session.access_token,
            'refresh_token': response.session.refresh_token,
            'user': {'email': user.email, 'id': user.id}
        }, status=status.HTTP_200_OK)

# resetpass_url ='http://localhost:3000/reset-password'
class PasswordResetView(APIView):
  
    def post(self, request):
        email = request.data.get('email')
        if not email:
            return Response({'error': 'Email is required'}, status=status.HTTP_400_BAD_REQUEST)

        # Validate email format
        import re
        email_pattern = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
        if not re.match(email_pattern, email):
            return Response({'error': 'Please enter a valid email address.'}, status=status.HTTP_400_BAD_REQUEST)

        try:
            # Trigger password reset email through Supabase
            # Supabase handles non-existent users gracefully and returns success for security
            response = supabase.auth.reset_password_for_email(
                email, 
                options={'redirect_to': settings.BASE_URL_RESET_PASSWORD}
            )
            
            # Always return success to prevent email enumeration attacks
            # Supabase will only send emails to registered users
            return Response({
                'message': 'If an account with this email exists, a password reset link has been sent.'
            }, status=status.HTTP_200_OK)
                
        except Exception as e:
            print(f"Password reset error: {str(e)}")
            return Response({
                'error': 'Failed to send password reset email. Please try again later.'
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


class PasswordResetConfirmView(APIView):
    def post(self, request):
        token = request.data.get('token')
        new_password = request.data.get('new_password')
        if not token or not new_password:
            return Response({'error': 'Token and new password are required.'}, status=status.HTTP_400_BAD_REQUEST)

        import requests
        from django.conf import settings

        project_url = getattr(settings, 'SUPABASE_URL', '').rstrip('/')
        api_key = getattr(settings, 'SUPABASE_SERVICE_ROLE_KEY', '') or getattr(settings, 'SUPABASE_ANON_KEY', '')
        if not project_url or not api_key:
            return Response({'error': 'Supabase configuration missing.'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

        url = f'{project_url}/auth/v1/user'
        headers = {
            'Authorization': f'Bearer {token}',
            'apikey': api_key,
            'Content-Type': 'application/json'
        }
        payload = {'password': new_password}

        try:
            resp = requests.put(url, json=payload, headers=headers)
            if resp.status_code != 200:
                return Response({'error': resp.json().get('message', 'Failed to update password')}, status=resp.status_code)
        except Exception as e:
            return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)

        return Response({'message': 'Password has been reset successfully.'}, status=status.HTTP_200_OK)

from gotrue.errors import AuthApiError

class VerifyOtpView(APIView):

    def post(self, request):
        access_token = request.data.get('access_token')
        if not access_token:
            return Response({'error': 'Access token is required'}, status=status.HTTP_400_BAD_REQUEST)

        try:
            user_response = supabase.auth.get_user(access_token)
        except AuthApiError as e:
            return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)

        user = user_response.user
        if not user:
            return Response({'error': 'Invalid token or user not found'}, status=status.HTTP_400_BAD_REQUEST)

        if user.user_metadata.get('email_confirmed'):
            return Response({'message': 'Email verified'}, status=status.HTTP_200_OK)
        else:
            return Response({'message': 'Email not verified'}, status=status.HTTP_400_BAD_REQUEST)

import uuid


class SupabaseOptions:
    def __init__(self, token):
        self.headers = {"Authorization": f"Bearer {token}"}
        self.auto_refresh_token = False
        self.persist_session = False
        self.detect_session_in_url = False
        self.storage = None
        self.flow_type = None
        self.httpx_client = None


class FileUploadView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request, *args, **kwargs):
        user = request.user
        folder = f"user_{user.id}"

        supabase = create_client(settings.SUPABASE_URL, settings.SUPABASE_SERVICE_ROLE_KEY)

        uploaded_file = request.FILES.get("file")
        if not uploaded_file:
            return Response({"error": "No file provided"}, status=status.HTTP_400_BAD_REQUEST)

        ext = uploaded_file.name.split(".")[-1].lower()
        
        # Validate file type
        if ext not in ['pdf', 'docx']:
            return Response(
                {"error": "Unsupported file type. Only PDF and DOCX files are supported."}, 
                status=status.HTTP_400_BAD_REQUEST
            )
        
        # Use original filename for storage
        import re
        original_filename = uploaded_file.name
        # Sanitize filename to remove unwanted characters and avoid path traversal
        sanitized_filename = re.sub(r'[^a-zA-Z0-9_.-]', '_', original_filename)
        storage_key = f"{folder}/{sanitized_filename}"

        try:
            file_bytes = uploaded_file.read()

            # 1) Upload to Supabase
            res = supabase.storage.from_("user-uploads").upload(
                storage_key,
                file_bytes,
                {"content-type": uploaded_file.content_type}
            )
            if isinstance(res, dict) and res.get("error"):
                return Response({"error": res["error"]["message"]}, status=status.HTTP_400_BAD_REQUEST)

            # 2) Save DB row & run embedding pipeline
            with transaction.atomic():
                collection_name = get_chroma_collection_name(user)

                # Save local copy so process_and_store_file can read from disk
                file_obj = File.objects.create(
                    user=user,
                    file=uploaded_file,          # FileField writes to MEDIA_ROOT/uploads/...
                    filename=uploaded_file.name, # original name
                    chroma_collection=collection_name,
                    storage_key=storage_key
                )

            # Path on local disk (since we just saved FileField)
            file_path = file_obj.file.path

            # 3) Create embeddings (sync)
            try:
                print(f"Starting embedding process for file: {file_path}")
                print(f"DEBUG: Database file ID = {file_obj.id}")
                print(f"DEBUG: Collection name = {collection_name}")
                print(f"DEBUG: User ID = {user.id}")
                
                process_and_store_file(user, file_path, collection_name, file_id=file_obj.id)
                print(f"Embedding process completed successfully for file: {file_path}")
                
                # Verify embeddings were created
                user_chroma_dir = get_user_chroma_dir(user)
                if Path(user_chroma_dir).exists():
                    try:
                        vs = Chroma(
                            persist_directory=user_chroma_dir,
                            embedding_function=text_emb,
                            collection_name=collection_name,
                        )
                        all_docs = vs.get()
                        if all_docs and 'documents' in all_docs:
                            print(f"✅ Verification: {len(all_docs['documents'])} embeddings found in ChromaDB")
                            # Check if our file_id is in the metadata
                            file_ids_in_chroma = set()
                            if 'metadatas' in all_docs:
                                for metadata in all_docs['metadatas']:
                                    if metadata and 'file_id' in metadata:
                                        file_ids_in_chroma.add(metadata['file_id'])
                            print(f"✅ File IDs in ChromaDB: {file_ids_in_chroma}")
                            if str(file_obj.id) in file_ids_in_chroma:
                                print(f"✅ File ID {file_obj.id} successfully stored in ChromaDB")
                            else:
                                print(f"⚠️ File ID {file_obj.id} not found in ChromaDB metadata")
                        else:
                            print(f"⚠️ No documents found in ChromaDB after upload")
                    except Exception as verify_error:
                        print(f"⚠️ Error verifying ChromaDB: {verify_error}")
                else:
                    print(f"⚠️ ChromaDB directory does not exist: {user_chroma_dir}")
                    
            except Exception as e:
                print(f"Embedding process failed for file {file_path}: {str(e)}")
                # Delete the file from database if embedding failed
                file_obj.delete()
                raise e

            # Delete local file after processing
            if os.path.exists(file_path):
                os.remove(file_path)

            return Response(
                {
                    "message": "File uploaded & embedded",
                    "file_id": file_obj.id,
                    "storage_key": storage_key
                },
                status=status.HTTP_201_CREATED
            )

        except Exception as e:
            # Optional: clean up supabase object if DB/embedding failed
            try:
                supabase.storage.from_("user-uploads").remove([storage_key])
            except Exception:
                pass

            return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

from django.conf import settings

from .serializers import UserSerializer
from rest_framework import generics
from rest_framework.permissions import IsAuthenticated

class UserProfileView(generics.RetrieveUpdateAPIView):
    serializer_class = UserSerializer
    permission_classes = [IsAuthenticated]

    def get_object(self):
        return self.request.user


class UserFilesView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request, *args, **kwargs):
        supabase = create_client(settings.SUPABASE_URL, settings.SUPABASE_SERVICE_ROLE_KEY)

        folder_name = f"user_{request.user.id}"

        try:
            files = supabase.storage.from_("user-uploads").list(
                path=folder_name,
                options={"limit": 100, "offset": 0, "sortBy": {"column": "name", "order": "asc"}},
            )

            # helper to build path once
            def full_path(name: str) -> str:
                return f"{folder_name}/{name}"

            # ----- OPTION A: bucket is PUBLIC -----
            # get_public_url = supabase.storage.from_("user-uploads").get_public_url

            # ----- OPTION B: bucket is PRIVATE (recommended) -----
            create_signed_url = supabase.storage.from_("user-uploads").create_signed_url
            EXPIRES_IN = 3600  # seconds

            files_info = []
            for f in files:
                if not f["name"].lower().endswith((".pdf", ".docx")):
                    continue

                path = full_path(f["name"])

                # PUBLIC:
                # url = get_public_url(path)

                # PRIVATE (signed):
                signed = create_signed_url(path, EXPIRES_IN)
                url = signed.get("signedURL") or signed.get("signed_url")  # lib versions differ

                files_info.append({
                    "name": f["name"],
                    "updated_at": f.get("updated_at"),
                    "created_at": f.get("created_at"),
                    "id": f.get("id"),
                    "url": url,
                })

            return Response({"files": files_info}, status=status.HTTP_200_OK)

        except Exception as e:
            return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


class UserFileDeleteView(APIView):
    permission_classes = [IsAuthenticated]

    # def delete(self, request, file_id):
    #     user = request.user
    #     supabase_token = request.headers.get("Authorization", "").replace("Bearer ", "")

    #     try:
    #         # Initialize supabase client with service role key
    #         supabase = create_client(
    #             settings.SUPABASE_URL,
    #             settings.SUPABASE_SERVICE_ROLE_KEY,
    #             options={
    #                 "auto_refresh_token": True,
    #                 "persist_session": True,
    #                 "detect_session_in_url": True,
    #                 "storage": None,
    #             },
    #         )

    #         # Validate file ownership by id
    #         file_obj = File.objects.filter(id=file_id, user=user).first()
    #         if not file_obj:
    #             return Response({"error": "File not found or you do not have permission."}, status=status.HTTP_404_NOT_FOUND)

    #         file_path = file_obj.storage_key

    #         # Remove file from Supabase storage
    #         removed = supabase.storage.from_("user-uploads").remove([file_path])
    #         if isinstance(removed, dict) and removed.get("error"):
    #             return Response({"error": removed["error"]["message"]}, status=status.HTTP_400_BAD_REQUEST)

    #         # Remove corresponding embeddings from ChromaDB
    #         try:
    #             collection = chroma_client.get_collection(name="user_files")  # Change to your collection name
    #             collection.delete(where={"file_id": str(file_id)})
    #         except Exception:
    #             pass  # Log error if needed

    #         # Remove DB record for the file
    #         file_obj.delete()

    #         return Response({"message": "File deleted."}, status=status.HTTP_200_OK)

    #     except Exception as e:
    #         return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    # permission_classes = [IsAuthenticated]

    def delete(self, request, file_name: str, *args, **kwargs):
        # Server-side key (or a valid Supabase user JWT in another header if you insist)
        supabase = create_client(settings.SUPABASE_URL, settings.SUPABASE_SERVICE_ROLE_KEY)

        user = request.user
        folder = f"user_{user.id}"
        file_path = f"{folder}/{file_name}"

        # (Optional) simple path traversal guard
        if "/" in file_name or ".." in file_name:
            return Response({"error": "Invalid file name."}, status=status.HTTP_400_BAD_REQUEST)

        try:
            # (Optional) verify ownership before delete
            # Cheap check using list; for large folders use your DB instead
            files = supabase.storage.from_("user-uploads").list(folder)
            if not any(f["name"] == file_name for f in files):
                return Response({"error": "File not found."}, status=status.HTTP_404_NOT_FOUND)

            # Remove from storage
            removed = supabase.storage.from_("user-uploads").remove([file_path])
            # v2: returns list of dicts; v1: dict with data/error
            if isinstance(removed, dict) and removed.get("error"):
                return Response({"error": removed["error"]["message"]}, status=status.HTTP_400_BAD_REQUEST)

            # Remove from DB (wrap to keep consistency)
            with transaction.atomic():
                from .models import File
                file_obj = File.objects.filter(user=user, storage_key=file_path).first()
                if file_obj:
                    if file_obj.chroma_collection:
                        try:
                            from langchain_chroma import Chroma
                        except ImportError:
                            from langchain_community.vectorstores import Chroma
                        import logging
                        
                        user_chroma_dir = get_user_chroma_dir(user.id)
                        
                        # Check if ChromaDB directory exists
                        if Path(user_chroma_dir).exists():
                            vs = Chroma(
                                collection_name=file_obj.chroma_collection,
                                embedding_function=text_emb,
                                persist_directory=user_chroma_dir,
                            )
                            
                            # Delete embeddings for this specific file
                            try:
                                print(f"Attempting to delete embeddings for file: {file_obj.filename} (ID: {file_obj.id})")
                                
                                # Get all documents to see what we're working with
                                all_docs = vs.get()
                                print(f"Total documents in ChromaDB: {len(all_docs['documents']) if all_docs and 'documents' in all_docs else 0}")
                                
                                # Debug: Show all metadata to understand the structure
                                if all_docs and 'metadatas' in all_docs:
                                    print("ChromaDB metadata structure:")
                                    for i, metadata in enumerate(all_docs['metadatas']):
                                        print(f"  Doc {i}: {metadata}")
                                
                                # Try multiple deletion strategies
                                try:
                                    deleted_count = 0
                                    
                                    # Strategy 1: Delete by file_id (string)
                                    try:
                                        vs.delete(where={"file_id": str(file_obj.id)})
                                        print(f"✓ Deleted by file_id (string): {file_obj.id}")
                                        deleted_count += 1
                                    except Exception as e1:
                                        print(f"✗ Failed to delete by file_id (string): {e1}")
                                    
                                    # Strategy 2: Delete by file_id (integer)
                                    try:
                                        vs.delete(where={"file_id": file_obj.id})
                                        print(f"✓ Deleted by file_id (integer): {file_obj.id}")
                                        deleted_count += 1
                                    except Exception as e2:
                                        print(f"✗ Failed to delete by file_id (integer): {e2}")
                                    
                                    # Strategy 3: Delete by filename
                                    try:
                                        vs.delete(where={"filename": file_obj.filename})
                                        print(f"✓ Deleted by filename: {file_obj.filename}")
                                        deleted_count += 1
                                    except Exception as e3:
                                        print(f"✗ Failed to delete by filename: {e3}")
                                    
                                    # Strategy 4: Delete by storage_key
                                    try:
                                        vs.delete(where={"storage_key": file_obj.storage_key})
                                        print(f"✓ Deleted by storage_key: {file_obj.storage_key}")
                                        deleted_count += 1
                                    except Exception as e4:
                                        print(f"✗ Failed to delete by storage_key: {e4}")
                                    
                                    # Strategy 5: Manual deletion by finding matching documents
                                    try:
                                        all_docs = vs.get()
                                        if all_docs and 'documents' in all_docs and 'metadatas' in all_docs and 'ids' in all_docs:
                                            matching_ids = []
                                            for i, metadata in enumerate(all_docs['metadatas']):
                                                if metadata:
                                                    file_id_match = str(metadata.get('file_id', '')) == str(file_obj.id)
                                                    filename_match = metadata.get('filename', '') == file_obj.filename
                                                    storage_key_match = metadata.get('storage_key', '') == file_obj.storage_key
                                                    
                                                    if file_id_match or filename_match or storage_key_match:
                                                        matching_ids.append(all_docs['ids'][i])
                                                        print(f"Found matching document: ID={all_docs['ids'][i]}, metadata={metadata}")
                                            
                                            if matching_ids:
                                                vs.delete(ids=matching_ids)
                                                print(f"✓ Deleted {len(matching_ids)} documents by ID matching")
                                                deleted_count += len(matching_ids)
                                            else:
                                                print("No matching documents found for manual deletion")
                                    except Exception as e5:
                                        print(f"✗ Failed to delete by manual matching: {e5}")
                                    
                                    # Strategy 6: Force delete all documents if no specific matches found
                                    if deleted_count == 0:
                                        try:
                                            print("No specific matches found, attempting to delete all documents...")
                                            all_docs = vs.get()
                                            if all_docs and 'ids' in all_docs and all_docs['ids']:
                                                vs.delete(ids=all_docs['ids'])
                                                print(f"✓ Deleted all {len(all_docs['ids'])} documents from ChromaDB")
                                                deleted_count = len(all_docs['ids'])
                                        except Exception as e6:
                                            print(f"✗ Failed to delete all documents: {e6}")
                                    
                                    if deleted_count > 0:
                                        print(f"✅ Successfully deleted {deleted_count} embeddings")
                                    else:
                                        print(f"⚠️ No embeddings were deleted for file: {file_obj.filename}")
                                    
                                    # Clean ChromaDB SQLite database
                                    print("🧹 Cleaning ChromaDB SQLite database...")
                                    clean_chroma_sqlite(user_chroma_dir, file_obj.chroma_collection)
                                    
                                except Exception as e:
                                    logging.error(f"Failed to delete embeddings for file_id {file_obj.id}: {e}")
                                    print(f"❌ Error during ChromaDB deletion: {e}")
                            except Exception as e:
                                    logging.error(f"Failed to initialize ChromaDB for file_id {file_obj.id}: {e}")
                                    print(f"❌ Error initializing ChromaDB: {e}")
                    else:       
                        print(f"ChromaDB directory does not exist for user {user.id}")
                    # Delete local file
                    file_obj.file.delete(save=False)
                    file_obj.delete()

            return Response({"message": "File deleted."}, status=status.HTTP_200_OK)
        except Exception as e:
            # log.exception("Supabase delete failed")
            return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

 

from rest_framework.generics import DestroyAPIView

class FileDeleteView(DestroyAPIView):
    permission_classes = [IsAuthenticated]
    queryset = File.objects.all()
    serializer_class = FileSerializer

    def perform_destroy(self, instance):
        # Delete file from Supabase Storage
        try:
            if getattr(instance, 'storage_key', None):
                supabase.storage.from_('user-uploads').remove([instance.storage_key])
        except Exception as e:
            pass  # Log error if needed
        # Delete from local storage and database
        instance.file.delete(save=False)
        instance.delete()


class FileListView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request):
        user_id = self.request.user.id
        files = File.objects.filter(user_id=user_id)
        serializer = FileSerializer(files, many=True)
        return Response(serializer.data)

class FileDeleteView(APIView):
    permission_classes = [IsAuthenticated]

    def delete(self, request, pk):
        user_id = self.request.user.id
        try:
            file_obj = File.objects.get(pk=pk, user_id=user_id)
        except File.DoesNotExist:
            return Response({'error': 'File not found.'}, status=404)
        # Remove from ChromaDB using LangChain Chroma vector store
        if file_obj.chroma_collection:
            try:
                from langchain_chroma import Chroma
            except ImportError:
                from langchain_community.vectorstores import Chroma
            import logging
            
            user_chroma_dir = get_user_chroma_dir(user_id)
            
            # Check if ChromaDB directory exists
            if Path(user_chroma_dir).exists():
                vs = Chroma(
                    collection_name=file_obj.chroma_collection,
                    embedding_function=text_emb,
                    persist_directory=user_chroma_dir,
                )
                
                # Delete embeddings for this specific file
                try:
                    print(f"Attempting to delete embeddings for file: {file_obj.filename} (ID: {file_obj.id})")
                    
                    # Get all documents to see what we're working with
                    all_docs = vs.get()
                    print(f"Total documents in ChromaDB: {len(all_docs['documents']) if all_docs and 'documents' in all_docs else 0}")
                    
                    # Debug: Show all metadata to understand the structure
                    if all_docs and 'metadatas' in all_docs:
                        print("ChromaDB metadata structure:")
                        for i, metadata in enumerate(all_docs['metadatas']):
                            print(f"  Doc {i}: {metadata}")
                    
                    # Try multiple deletion strategies
                    try:
                        deleted_count = 0
                        
                        # Strategy 1: Delete by file_id (string)
                        try:
                            vs.delete(where={"file_id": str(file_obj.id)})
                            print(f"✓ Deleted by file_id (string): {file_obj.id}")
                            deleted_count += 1
                        except Exception as e1:
                            print(f"✗ Failed to delete by file_id (string): {e1}")
                        
                        # Strategy 2: Delete by file_id (integer)
                        try:
                            vs.delete(where={"file_id": file_obj.id})
                            print(f"✓ Deleted by file_id (integer): {file_obj.id}")
                            deleted_count += 1
                        except Exception as e2:
                            print(f"✗ Failed to delete by file_id (integer): {e2}")
                        
                        # Strategy 3: Delete by filename
                        try:
                            vs.delete(where={"filename": file_obj.filename})
                            print(f"✓ Deleted by filename: {file_obj.filename}")
                            deleted_count += 1
                        except Exception as e3:
                            print(f"✗ Failed to delete by filename: {e3}")
                        
                        # Strategy 4: Delete by storage_key
                        try:
                            vs.delete(where={"storage_key": file_obj.storage_key})
                            print(f"✓ Deleted by storage_key: {file_obj.storage_key}")
                            deleted_count += 1
                        except Exception as e4:
                            print(f"✗ Failed to delete by storage_key: {e4}")
                        
                        # Strategy 5: Manual deletion by finding matching documents
                        try:
                            all_docs = vs.get()
                            if all_docs and 'documents' in all_docs and 'metadatas' in all_docs and 'ids' in all_docs:
                                matching_ids = []
                                for i, metadata in enumerate(all_docs['metadatas']):
                                    if metadata:
                                        file_id_match = str(metadata.get('file_id', '')) == str(file_obj.id)
                                        filename_match = metadata.get('filename', '') == file_obj.filename
                                        storage_key_match = metadata.get('storage_key', '') == file_obj.storage_key
                                        
                                        if file_id_match or filename_match or storage_key_match:
                                            matching_ids.append(all_docs['ids'][i])
                                            print(f"Found matching document: ID={all_docs['ids'][i]}, metadata={metadata}")
                                
                                if matching_ids:
                                    vs.delete(ids=matching_ids)
                                    print(f"✓ Deleted {len(matching_ids)} documents by ID matching")
                                    deleted_count += len(matching_ids)
                                else:
                                    print("No matching documents found for manual deletion")
                        except Exception as e5:
                            print(f"✗ Failed to delete by manual matching: {e5}")
                        
                        # Strategy 6: Force delete all documents if no specific matches found
                        if deleted_count == 0:
                            try:
                                print("No specific matches found, attempting to delete all documents...")
                                all_docs = vs.get()
                                if all_docs and 'ids' in all_docs and all_docs['ids']:
                                    vs.delete(ids=all_docs['ids'])
                                    print(f"✓ Deleted all {len(all_docs['ids'])} documents from ChromaDB")
                                    deleted_count = len(all_docs['ids'])
                            except Exception as e6:
                                print(f"✗ Failed to delete all documents: {e6}")
                        
                        if deleted_count > 0:
                            print(f"✅ Successfully deleted {deleted_count} embeddings")
                        else:
                            print(f"⚠️ No embeddings were deleted for file: {file_obj.filename}")
                        
                        # Clean ChromaDB SQLite database
                        print("🧹 Cleaning ChromaDB SQLite database...")
                        clean_chroma_sqlite(user_chroma_dir, file_obj.chroma_collection)
                        
                    except Exception as e:
                        logging.error(f"Failed to delete embeddings for file_id {file_obj.id}: {e}")
                        print(f"❌ Error during ChromaDB deletion: {e}")
                except Exception as e:
                    logging.error(f"Failed to initialize ChromaDB for file_id {file_obj.id}: {e}")
                    print(f"❌ Error initializing ChromaDB: {e}")
            else:
                print(f"ChromaDB directory does not exist for user {user_id}")
        # Delete file from storage and DB
        file_obj.file.delete(save=False)
        file_obj.delete()
        return Response({'message': 'File deleted.'}, status=204)

class ChatListCreateView(generics.ListCreateAPIView):
    serializer_class = ChatSerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        user = self.request.user  # this is a Django User instance now
        user_id = user.id
        # Sort chats by the latest updated timestamp (most recent first)
        return Chat.objects.filter(user_id=user_id).order_by('-updated_at', '-created_at')

    def perform_create(self, serializer):
        user = self.request.user  # this is a Django User instance now
        user_id = user.id
        serializer.save(user_id=user_id)


class ChatUpdateDeleteView(APIView):
    permission_classes = [IsAuthenticated]

    def get_object(self, chat_id, user):
        return get_object_or_404(Chat, id=chat_id, user_id=user.id)

    def patch(self, request, chat_id):
        chat = self.get_object(chat_id, request.user)
        title = request.data.get('title')
        if not title or not title.strip():
            return Response({'error': 'Title is required.'}, status=status.HTTP_400_BAD_REQUEST)
        
        # Update title and explicitly update the updated_at field
        chat.title = title.strip()
        chat.updated_at = timezone.now()
        chat.save()
        
        serializer = ChatSerializer(chat)
        return Response(serializer.data, status=status.HTTP_200_OK)

    def delete(self, request, chat_id):
        chat = self.get_object(chat_id, request.user)
        chat.delete()
        return Response({'message': 'Chat deleted.'}, status=status.HTTP_204_NO_CONTENT)


def serialize_message(m: Message):
    return {
        "id": m.id,
        "chat_id": str(m.chat_id),
        "sender_id": str(m.sender_id) if m.sender_id else None,
        "sender_type": m.sender_type,
        "content": m.content,
        "timestamp": m.timestamp.isoformat(),
    }
    
class MessageListCreateView(generics.ListCreateAPIView):
    serializer_class = MessageSerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        chat_id = self.kwargs["chat_id"]
        return (Message.objects
                .filter(chat_id=chat_id, chat__user_id=self.request.user.id)
                .order_by("timestamp"))

    def list(self, request, *args, **kwargs):
        msgs = self.get_queryset()
        data = [{
            "id": m.id,
            "chat_id": str(m.chat_id),
            "user_message": m.content if m.sender_type == "user" else None,
            "ai_response":  m.content if m.sender_type == "ai"   else None,
            "rag_sources":  m.sources  if m.sender_type == "ai"   else None,
        } for m in msgs]
        return Response({"messages": data}, status=status.HTTP_200_OK)

    def create(self, request, *args, **kwargs):
        try:
            user    = request.user
            chat_id = self.kwargs["chat_id"]
            chat    = get_object_or_404(Chat, id=chat_id, user_id=user.id)

            # 1) Save user message
            ser = self.get_serializer(data=request.data)
            ser.is_valid(raise_exception=True)
            user_msg = ser.save(sender=user, sender_type="user", chat=chat)

            # 2) RAG -> AI message
            rag = run_rag_pipeline(user, user_msg.content, chat_id=chat_id)
            ai_msg = Message.objects.create(
                chat=chat,
                sender=None,
                sender_type="ai",
                content=rag["answer"],
                sources=rag.get("sources", [])
            )

            # 3) Update chat's updated_at field to reflect new activity
            chat.updated_at = timezone.now()
            chat.save()

            # 4) Flat response
            return Response(
                {
                    "id": ai_msg.id,                     # or user_msg.id / chat_id — your choice
                    "chat_id": str(chat.id),
                    "user_message": user_msg.content,
                    "ai_response":  ai_msg.content,
                    "rag_sources": rag.get("sources", [])  # optional - now includes filename, page, and content
                },
                status=status.HTTP_201_CREATED
            )
        except Exception as e:
            print(f"Error in message creation: {str(e)}")
            return Response(
                {"error": "Failed to process message. Please try again."}, 
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )

from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status
from supabase import create_client
from django.conf import settings

class SendVerificationEmailView(APIView):
    """
    POST endpoint to send a verification email to a user if not verified in Supabase.
    Accepts 'email' in POST data (optional if authenticated user).
    """
    def post(self, request):
        email = request.data.get('email')
        if not email and request.user and hasattr(request.user, 'email'):
            email = request.user.email
        if not email:
            return Response({'error': 'Email is required.'}, status=status.HTTP_400_BAD_REQUEST)
        supabase = create_client(settings.SUPABASE_URL, settings.SUPABASE_SERVICE_ROLE_KEY)
        # Check user status in Supabase
        try:
            # list_users() does not accept email param; fetch all and filter
            users = supabase.auth.admin.list_users() or []
            user_obj = None
            for u in users:
                # Support both dict and object for user
                email_val = u['email'] if isinstance(u, dict) else getattr(u, 'email', None)
                if email_val and email_val.lower() == email.lower():
                    user_obj = u
                    break
            if not user_obj:
                return Response({'error': 'User not found.'}, status=status.HTTP_404_NOT_FOUND)
            # Optionally, get latest user details by id
            user_id = user_obj['id'] if isinstance(user_obj, dict) else getattr(user_obj, 'id', None)
            user_details = supabase.auth.admin.get_user_by_id(user_id)
            # user_details may be dict or object
            email_confirmed_at = None
            if isinstance(user_details, dict):
                data = user_details.get('data', {})
                email_confirmed_at = data.get('email_confirmed_at')
            else:
                email_confirmed_at = getattr(user_details, 'email_confirmed_at', None)
            if email_confirmed_at:
                return Response({'error': 'User already verified.'}, status=status.HTTP_400_BAD_REQUEST)
            # Resend verification email for existing unverified user
          
            try:
                # print("email",email)
                # print("BASE_URL_SIGNIN",settings.BASE_URL_SIGNIN)
                resend_resp = supabase.auth.resend({
                    'type': 'signup',
                    'email': email,
                    'options': {
                        'email_redirect_to': settings.BASE_URL_SIGNIN
                    }
                })
                if resend_resp is not None:
                    # Check for error or success in resend_resp
                    error = None
                    if isinstance(resend_resp, dict):
                        error = resend_resp.get('error')
                    else:
                        error = getattr(resend_resp, 'error', None)
                    if error:
                        return Response({'error': f'Verification email not sent: {error}'}, status=status.HTTP_400_BAD_REQUEST)
                    return Response(
                        {'message': 'Verification email has been resent. Please check your email and verify your account.'},
                        status=status.HTTP_200_OK
                    )
                else:
                    return Response({'error': 'Failed to resend verification email'}, status=status.HTTP_400_BAD_REQUEST)
            except Exception as resend_error:
                return Response(
                    {'error': f'Failed to resend verification email: {str(resend_error)}'},
                    status=status.HTTP_500_INTERNAL_SERVER_ERROR
                )
        except Exception as e:
            return Response({'error': f'Failed to send verification email: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


class TokenRefreshView(APIView):
    def post(self, request):
        refresh_token = request.data.get("refresh_token")
        if not refresh_token:
            return Response({"error": "refresh_token is required"}, status=status.HTTP_400_BAD_REQUEST)

        supabase_url = settings.SUPABASE_URL
        supabase_key = settings.SUPABASE_KEY
        supabase = create_client(supabase_url, supabase_key)

        try:
            data = supabase.auth.refresh_session(refresh_token)
            if data.session and data.session.access_token and data.session.refresh_token:
                return Response({
                    "access_token": data.session.access_token,
                    "refresh_token": data.session.refresh_token
                })
            else:
                return Response({"error": "Failed to refresh token"}, status=status.HTTP_400_BAD_REQUEST)
        except Exception as e:
            return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

def get_user_files_from_chroma_sqlite(user):
    """
    Directly query ChromaDB SQLite database to get list of files for a user
    Returns a list of file information from the database
    """
    try:
        user_chroma_dir = get_user_chroma_dir(user)
        sqlite_path = os.path.join(user_chroma_dir, "chroma.sqlite3")
        
        if not os.path.exists(sqlite_path):
            print(f"ChromaDB SQLite file not found: {sqlite_path}")
            return []
            
        print(f"Querying ChromaDB SQLite: {sqlite_path}")
        
        # Connect to SQLite database
        conn = sqlite3.connect(sqlite_path)
        cursor = conn.cursor()
        
        # Get all tables
        cursor.execute("SELECT name FROM sqlite_master WHERE type='table';")
        tables = cursor.fetchall()
        print(f"Found tables: {[table[0] for table in tables]}")
        
        # Get collection name for this user
        collection_name = get_chroma_collection_name(user)
        print(f"Looking for collection: {collection_name}")
        
        # Get segment_id for the collection
        cursor.execute("SELECT id FROM segments WHERE name = ?;", (collection_name,))
        segment_result = cursor.fetchone()
        
        if not segment_result:
            print(f"No segment found for collection: {collection_name}")
            conn.close()
            return []
            
        segment_id = segment_result[0]
        print(f"Found segment_id: {segment_id} for collection: {collection_name}")
        
        # Get unique file_ids from embedding_metadata table
        cursor.execute("""
            SELECT DISTINCT em.value 
            FROM embedding_metadata em 
            WHERE em.segment_id = ? AND em.key = 'file_id'
        """, (segment_id,))
        
        file_ids = [row[0] for row in cursor.fetchall()]
        print(f"Found file IDs in ChromaDB: {file_ids}")
        
        conn.close()
        
        if not file_ids:
            return []
        
        # Get file information from Django database
        try:
            file_ids_int = [int(fid) for fid in file_ids]
            user_files = File.objects.filter(
                user_id=user.id, 
                id__in=file_ids_int
            ).order_by('uploaded_at')
            
            file_list = []
            for file in user_files:
                file_info = {
                    "filename": file.filename,
                    "uploaded_at": file.uploaded_at.strftime("%Y-%m-%d %H:%M:%S"),
                    "file_type": file.filename.split('.')[-1].upper() if '.' in file.filename else "Unknown",
                    "file_id": file.id
                }
                file_list.append(file_info)
                print(f"✓ Found file in ChromaDB: {file.filename} (ID: {file.id})")
            
            return file_list
            
        except ValueError as e:
            print(f"Error converting file IDs to integers: {e}")
            return []
            
    except Exception as e:
        print(f"Error querying ChromaDB SQLite: {e}")
        return []

def get_user_files_from_chroma(user):
    """
    Get list of files for a user from ChromaDB
    Returns a list of file information for files that have been processed and stored in ChromaDB
    """
    try:
        print(f"🔍 Getting files from ChromaDB for user {user.id}")
        print(f"🔍 User details: ID={user.id}, Email={getattr(user, 'email', 'N/A')}")
        
        # Get user's ChromaDB directory and collection
        user_chroma_dir = get_user_chroma_dir(user)
        collection_name = get_chroma_collection_name(user)
        
        print(f"🔍 ChromaDB directory: {user_chroma_dir}")
        print(f"🔍 Collection name: {collection_name}")
        
        # Check if ChromaDB directory exists
        if not Path(user_chroma_dir).exists():
            print(f"❌ ChromaDB directory does not exist: {user_chroma_dir}")
            return []
        
        # Connect to ChromaDB
        try:
            vs = Chroma(
                persist_directory=user_chroma_dir,
                embedding_function=text_emb,
                collection_name=collection_name,
            )
        except Exception as e:
            print(f"❌ Error connecting to ChromaDB: {e}")
            print(f"🔍 Checking directory permissions for: {user_chroma_dir}")
            if os.path.exists(user_chroma_dir):
                import stat
                st = os.stat(user_chroma_dir)
                print(f"🔍 Directory permissions: {oct(st.st_mode)}")
                print(f"🔍 Directory owner: {st.st_uid}")
                print(f"🔍 Directory group: {st.st_gid}")
            raise e
        
        # Get all documents from ChromaDB
        all_docs = vs.get()
        print(f"🔍 Total documents in ChromaDB: {len(all_docs['documents']) if all_docs and 'documents' in all_docs else 0}")
        
        if not all_docs or 'documents' not in all_docs or 'metadatas' not in all_docs:
            print(f"❌ No documents found in ChromaDB for user {user.id}")
            return []
        
        # Extract unique file IDs from ChromaDB metadata
        file_ids = set()
        for metadata in all_docs['metadatas']:
            if metadata and 'file_id' in metadata:
                file_ids.add(metadata['file_id'])
                print(f"🔍 Found file_id in ChromaDB: {metadata['file_id']}")
        
        print(f"🔍 Unique file IDs found in ChromaDB: {list(file_ids)}")
        
        if not file_ids:
            print(f"❌ No file IDs found in ChromaDB metadata")
            return []
        
        # Get file information from Django database for files that exist in ChromaDB
        try:
            file_ids_int = [int(fid) for fid in file_ids]
            user_files = File.objects.filter(
                user_id=user.id, 
                id__in=file_ids_int
            ).order_by('uploaded_at')
            
            print(f"🔍 Found {user_files.count()} files in database that match ChromaDB file_ids")
            
            file_list = []
            for file in user_files:
                file_info = {
                    "filename": file.filename,
                    "uploaded_at": file.uploaded_at.strftime("%Y-%m-%d %H:%M:%S"),
                    "file_type": file.filename.split('.')[-1].upper() if '.' in file.filename else "Unknown",
                    "file_id": file.id
                }
                file_list.append(file_info)
                print(f"✓ Found file in ChromaDB: {file.filename} (ID: {file.id}, User ID: {file.user_id})")
            
            print(f"✅ Found {len(file_list)} files in ChromaDB for user {user.id}")
            return file_list
            
        except ValueError as e:
            print(f"❌ Error converting file IDs to integers: {e}")
            return []
            
    except Exception as e:
        print(f"❌ Error getting files from ChromaDB: {e}")
        import traceback
        traceback.print_exc()
        return []

def test_user_chroma_document_listing(user):
    """
    Test function to verify ChromaDB document listing for a specific user
    """
    print(f"🧪 Testing ChromaDB document listing for user {user.id}")
    
    # Test 1: Check if user exists
    print(f"🧪 User ID: {user.id}")
    print(f"🧪 User email: {getattr(user, 'email', 'N/A')}")
    
    # Test 2: Check ChromaDB directory
    user_chroma_dir = get_user_chroma_dir(user)
    collection_name = get_chroma_collection_name(user)
    print(f"🧪 ChromaDB directory: {user_chroma_dir}")
    print(f"🧪 Collection name: {collection_name}")
    print(f"🧪 Directory exists: {Path(user_chroma_dir).exists()}")
    
    # Test 3: Check total files in database
    total_files = File.objects.count()
    print(f"🧪 Total files in database: {total_files}")
    
    # Test 4: Check files for this user
    user_files = File.objects.filter(user_id=user.id)
    print(f"🧪 Files for user {user.id}: {user_files.count()}")
    
    # Test 5: List all files for this user
    for file in user_files:
        print(f"🧪 File: {file.filename} (ID: {file.id}, User: {file.user_id})")
    
    # Test 6: Call the actual ChromaDB function
    result = get_user_files_from_chroma(user)
    print(f"🧪 ChromaDB function result: {len(result)} files")
    
    return result

def check_and_fix_chroma_permissions(user_chroma_dir):
    """
    Check and fix ChromaDB directory permissions
    """
    try:
        import stat
        
        print(f"🔧 Checking ChromaDB permissions for: {user_chroma_dir}")
        
        if not os.path.exists(user_chroma_dir):
            print(f"🔧 Creating ChromaDB directory: {user_chroma_dir}")
            os.makedirs(user_chroma_dir, exist_ok=True)
        
        # Check current permissions
        st = os.stat(user_chroma_dir)
        current_permissions = oct(st.st_mode)
        current_owner = st.st_uid
        current_group = st.st_gid
        
        print(f"🔧 Current permissions: {current_permissions}")
        print(f"🔧 Current owner: {current_owner}")
        print(f"🔧 Current group: {current_group}")
        
        # Check if we can write to the directory
        test_file = os.path.join(user_chroma_dir, "test_write.tmp")
        try:
            with open(test_file, 'w') as f:
                f.write("test")
            os.remove(test_file)
            print(f"✅ Write permissions OK")
            return True
        except Exception as e:
            print(f"❌ Write permissions failed: {e}")
            
            # Try to fix permissions
            try:
                print(f"🔧 Attempting to fix permissions...")
                subprocess.run(['sudo', 'chown', '-R', 'chatapp:www-data', user_chroma_dir], check=True)
                subprocess.run(['sudo', 'chmod', '-R', '775', user_chroma_dir], check=True)
                print(f"✅ Permissions fixed")
                return True
            except Exception as fix_error:
                print(f"❌ Failed to fix permissions: {fix_error}")
                return False
                
    except Exception as e:
        print(f"❌ Error checking permissions: {e}")
        return False